from pathlib import Path
from typing import List, Optional

from .docx_utils import (
    TemporaryImageDir,
    analyze_images_to_sections,
    extract_images_from_docx,
    extract_xml_texts,
    table_to_markdown,
)


def _append_header_footer(result: List[str], doc) -> None:
    items = []
    for idx, section in enumerate(doc.sections, start=1):
        for label, part in (("页眉", section.header), ("页脚", section.footer)):
            texts = [p.text.strip() for p in part.paragraphs if p.text.strip()]
            if texts:
                items.append((idx, label, texts))
    if not items:
        return

    result.append("## 页眉页脚\n")
    for idx, label, texts in items:
        result.append(f"### 第 {idx} 节{label}")
        result.extend(texts)
        result.append("")


def _append_extra_xml_parts(result: List[str], file_path: str) -> None:
    parts = extract_xml_texts(
        file_path,
        [
            "word/comments.xml",
            "word/footnotes.xml",
            "word/endnotes.xml",
        ],
    )
    labels = {
        "word/comments.xml": "批注",
        "word/footnotes.xml": "脚注",
        "word/endnotes.xml": "尾注",
    }
    if not parts:
        return

    result.append("## 批注/脚注/尾注\n")
    for member, texts in parts.items():
        result.append(f"### {labels.get(member, member)}")
        result.append("".join(texts).strip())
        result.append("")


def read_docx(
    file_path: str,
    use_image_analysis: bool = True,
    max_images: Optional[int] = None,
    min_image_kb: float = 3.0,
    image_prompt: Optional[str] = None,
) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("缺少 python-docx 包，请安装 python-docx 或使用 Codex 自带运行时。") from exc

    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"文档不存在: {file_path}")
    if path.suffix.lower() != ".docx":
        raise ValueError(f"read_docx 只支持 .docx，当前文件: {path.suffix}")

    result: List[str] = [f"# 文档: {path.name}\n", "## 文档正文\n"]
    doc = Document(str(path))

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style_name = para.style.name if para.style else "Normal"
        if "Heading" in style_name or "heading" in style_name:
            level = style_name.replace("Heading", "").replace("heading", "").strip()
            prefix = "#" * (int(level) if level.isdigit() else 2)
            result.append(f"{prefix} {text}\n")
        elif "Title" in style_name:
            result.append(f"# {text}\n")
        else:
            result.append(f"{text}\n")

    if doc.tables:
        result.append("## 表格内容\n")
        for table_idx, table in enumerate(doc.tables, start=1):
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            result.append(f"### 表格 {table_idx}")
            result.append(table_to_markdown(rows))
            result.append("")

    _append_header_footer(result, doc)
    _append_extra_xml_parts(result, str(path))

    if use_image_analysis:
        with TemporaryImageDir("docx_images_") as temp_dir:
            images = extract_images_from_docx(str(path), temp_dir, min_size_kb=min_image_kb)
            result.extend(analyze_images_to_sections(images, max_images=max_images, prompt=image_prompt))

    return "\n".join(result)
